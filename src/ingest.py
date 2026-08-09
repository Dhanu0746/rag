
"""
Document ingestion pipeline.

Loads supported documents (PDF, DOCX, TXT, Markdown),
chunks them using LlamaIndex, creates embeddings,
stores them in ChromaDB, and builds a BM25 index.

DOCX files are parsed explicitly using python-docx so that
the underlying ZIP/binary contents of the .docx file are
never passed into the RAG pipeline.
"""

import argparse
import pickle
from pathlib import Path

import chromadb
from docx import Document as DocxDocument
from llama_index.core import Document, SimpleDirectoryReader
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from rank_bm25 import BM25Okapi


# ---------------------------------------------------
# Paths / Configuration
# ---------------------------------------------------

DATA_DIR = Path(__file__).parent.parent / "data"
STORE_DIR = Path(__file__).parent.parent / "storage"

EMBED_MODEL_NAME = "BAAI/bge-small-en-v1.5"


# ---------------------------------------------------
# Load Documents
# ---------------------------------------------------

def load_documents(username) -> list[Document]:
    """
    Load supported documents from data/<username>/.

    Supported:
        - PDF
        - DOCX
        - TXT
        - Markdown

    DOCX files are handled explicitly using python-docx.
    This avoids accidentally reading the internal ZIP/binary
    structure of a DOCX file as text.
    """

    user_data_dir = DATA_DIR / username

    if not user_data_dir.exists():
        print(
            f"[ERROR] Data directory does not exist: "
            f"{user_data_dir}"
        )
        return []

    documents = []

    # ---------------------------------------------------
    # Walk through all files
    # ---------------------------------------------------

    for file_path in user_data_dir.rglob("*"):

        if not file_path.is_file():
            continue

        extension = file_path.suffix.lower()

        # ---------------------------------------------------
        # DOCX
        # ---------------------------------------------------

        if extension == ".docx":

            try:
                docx_file = DocxDocument(str(file_path))

                paragraphs = []

                # Extract normal paragraphs
                for paragraph in docx_file.paragraphs:

                    text = paragraph.text.strip()

                    if text:
                        paragraphs.append(text)

                # Extract text from tables as well
                for table in docx_file.tables:

                    for row in table.rows:

                        row_text = []

                        for cell in row.cells:

                            cell_text = cell.text.strip()

                            if cell_text:
                                row_text.append(cell_text)

                        if row_text:
                            paragraphs.append(
                                " | ".join(row_text)
                            )

                text = "\n".join(paragraphs)

                if text.strip():

                    documents.append(
                        Document(
                            text=text,
                            metadata={
                                "file_name": file_path.name,
                                "source": str(file_path),
                                "file_type": "docx",
                            },
                        )
                    )

                    print(
                        f"[DOCX] Loaded: {file_path.name} "
                        f"({len(text)} characters)"
                    )

                else:

                    print(
                        f"[WARNING] DOCX contains no readable text: "
                        f"{file_path.name}"
                    )

            except Exception as e:

                print(
                    f"[ERROR] Failed to read DOCX "
                    f"{file_path.name}: {e}"
                )

        # ---------------------------------------------------
        # PDF / TXT / Markdown
        # ---------------------------------------------------

        elif extension in {".pdf", ".txt", ".md"}:

            try:

                reader = SimpleDirectoryReader(
                    input_files=[str(file_path)]
                )

                loaded_documents = reader.load_data()

                documents.extend(loaded_documents)

                print(
                    f"[{extension.upper().replace('.', '')}] "
                    f"Loaded: {file_path.name}"
                )

            except Exception as e:

                print(
                    f"[ERROR] Failed to read "
                    f"{file_path.name}: {e}"
                )

    print()
    print(f"Loaded {len(documents)} documents")

    return documents


# ---------------------------------------------------
# Chunk Documents
# ---------------------------------------------------

def chunk_documents(
    documents: list[Document],
    chunk_size: int,
    chunk_overlap: int,
):
    """
    Split documents into chunks using LlamaIndex SentenceSplitter.

    SentenceSplitter is used for all document types because the
    input has already been converted into clean text.
    """

    sentence_splitter = SentenceSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    final_nodes = []

    for doc in documents:

        try:

            nodes = sentence_splitter.get_nodes_from_documents(
                [doc]
            )

            final_nodes.extend(nodes)

        except Exception as e:

            print(
                f"[WARNING] Failed to chunk document: {e}"
            )

    print(
        f"Produced {len(final_nodes)} chunks"
    )

    return final_nodes


# ---------------------------------------------------
# Dense Index
# ---------------------------------------------------

def build_dense_index(nodes, config_name):
    """
    Create embeddings and store them in ChromaDB.

    The existing collection is deleted first so that old/stale
    chunks cannot remain after re-ingestion.
    """

    if not nodes:
        print("[ERROR] No nodes available for dense indexing.")
        return

    embed_model = HuggingFaceEmbedding(
        model_name=EMBED_MODEL_NAME
    )

    chroma_path = STORE_DIR / "chroma"

    chroma_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    chroma_client = chromadb.PersistentClient(
        path=str(chroma_path)
    )

    # ---------------------------------------------------
    # Remove old collection
    # ---------------------------------------------------

    try:

        chroma_client.delete_collection(
            name=config_name
        )

        print(
            f"[dense] Deleted existing Chroma collection "
            f"'{config_name}'"
        )

    except Exception:

        # Collection doesn't exist yet
        pass

    # ---------------------------------------------------
    # Create fresh collection
    # ---------------------------------------------------

    collection = chroma_client.create_collection(
        name=config_name
    )

    # ---------------------------------------------------
    # Prepare texts
    # ---------------------------------------------------

    texts = [
        node.get_content()
        for node in nodes
    ]

    print(
        f"[dense] Creating embeddings for "
        f"{len(texts)} chunks..."
    )

    embeddings = embed_model.get_text_embedding_batch(
        texts,
        show_progress=True,
    )

    ids = []
    metadatas = []

    # ---------------------------------------------------
    # Metadata
    # ---------------------------------------------------

    for i, node in enumerate(nodes):

        meta = dict(node.metadata)

        source = (
            meta.get("file_name")
            or meta.get("filename")
            or meta.get("source")
            or "Unknown"
        )

        page = (
            meta.get("page_label")
            or meta.get("page")
            or meta.get("page_number")
            or ""
        )

        ids.append(
            f"{config_name}_{i}"
        )

        metadatas.append(
            {
                "source": str(source),
                "page": str(page),
                "chunk_id": i,
                "text": node.get_content(),
            }
        )

    # ---------------------------------------------------
    # Store in Chroma
    # ---------------------------------------------------

    collection.add(
        ids=ids,
        embeddings=embeddings,
        documents=texts,
        metadatas=metadatas,
    )

    print(
        f"[dense] Indexed {len(nodes)} chunks "
        f"into Chroma collection '{config_name}'"
    )


# ---------------------------------------------------
# Sparse Index
# ---------------------------------------------------

def build_sparse_index(nodes, config_name):
    """
    Build and save a BM25 sparse retrieval index.
    """

    if not nodes:
        print("[ERROR] No nodes available for sparse indexing.")
        return

    # ---------------------------------------------------
    # Tokenize
    # ---------------------------------------------------

    tokenized = [
        node.get_content().lower().split()
        for node in nodes
    ]

    bm25 = BM25Okapi(tokenized)

    # ---------------------------------------------------
    # Ensure storage directory exists
    # ---------------------------------------------------

    STORE_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    texts = [
        node.get_content()
        for node in nodes
    ]

    metadata = []

    # ---------------------------------------------------
    # Metadata
    # ---------------------------------------------------

    for i, node in enumerate(nodes):

        meta = dict(node.metadata)

        source = (
            meta.get("file_name")
            or meta.get("filename")
            or meta.get("source")
            or "Unknown"
        )

        page = (
            meta.get("page_label")
            or meta.get("page")
            or meta.get("page_number")
            or ""
        )

        metadata.append(
            {
                "source": str(source),
                "page": str(page),
                "chunk_id": i,
            }
        )

    # ---------------------------------------------------
    # Save BM25
    # ---------------------------------------------------

    bm25_path = (
        STORE_DIR / f"bm25_{config_name}.pkl"
    )

    with open(
        bm25_path,
        "wb",
    ) as f:

        pickle.dump(
            {
                "bm25": bm25,
                "texts": texts,
                "metadatas": metadata,
            },
            f,
        )

    print(
        f"[sparse] Indexed {len(nodes)} chunks "
        f"into BM25 index"
    )

    print(
        f"[sparse] Saved BM25 index to "
        f"{bm25_path}"
    )


# ---------------------------------------------------
# Main
# ---------------------------------------------------

def main():

    parser = argparse.ArgumentParser(
        description="Ingest documents into the RAG indexes."
    )

    parser.add_argument(
        "--chunk-size",
        type=int,
        default=512,
        help="Maximum chunk size.",
    )

    parser.add_argument(
        "--chunk-overlap",
        type=int,
        default=64,
        help="Chunk overlap.",
    )

    parser.add_argument(
        "--config-name",
        default="default",
        help=(
            "Also used as the username/data subfolder "
            "(data/<config-name>/), matching how the "
            "Streamlit app names indexes."
        ),
    )

    args = parser.parse_args()

    # ---------------------------------------------------
    # Load
    # ---------------------------------------------------

    print()
    print("=" * 60)
    print("DOCUMENT INGESTION")
    print("=" * 60)
    print()

    documents = load_documents(
        args.config_name
    )

    if not documents:

        print(
            f"\n[ERROR] No documents found in "
            f"data/{args.config_name}/.\n"
        )

        print(
            "Add .pdf, .md, .txt, or .docx files "
            "there first, or upload through the "
            "Streamlit Document Manager."
        )

        return

    # ---------------------------------------------------
    # Chunk
    # ---------------------------------------------------

    nodes = chunk_documents(
        documents,
        args.chunk_size,
        args.chunk_overlap,
    )

    if not nodes:

        print(
            "[ERROR] No chunks were produced."
        )

        return

    # ---------------------------------------------------
    # Dense
    # ---------------------------------------------------

    build_dense_index(
        nodes,
        args.config_name,
    )

    # ---------------------------------------------------
    # Sparse
    # ---------------------------------------------------

    build_sparse_index(
        nodes,
        args.config_name,
    )

    # ---------------------------------------------------
    # Finished
    # ---------------------------------------------------

    print()
    print("=" * 60)
    print("INGESTION COMPLETE")
    print("=" * 60)
    print(
        f"Documents : {len(documents)}"
    )
    print(
        f"Chunks    : {len(nodes)}"
    )
    print(
        f"Collection: {args.config_name}"
    )
    print()


# ---------------------------------------------------
# Entry Point
# ---------------------------------------------------

if __name__ == "__main__":
    main()
